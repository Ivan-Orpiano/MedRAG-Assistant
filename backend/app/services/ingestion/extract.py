"""Text extraction for PDF / DOCX / TXT with per-page granularity and an OCR
fallback for scanned PDF pages (a page with an empty text layer is rasterized
with pdf2image and passed through Tesseract)."""
import io
from dataclasses import dataclass, field

from app.core.config import get_settings
from app.core.logging import get_logger

logger = get_logger(__name__)

SUPPORTED_MIME = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}


@dataclass
class ExtractedPage:
    page_number: int | None
    text: str
    used_ocr: bool = False


@dataclass
class ExtractionResult:
    pages: list[ExtractedPage] = field(default_factory=list)

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def ocr_pages(self) -> int:
        return sum(1 for p in self.pages if p.used_ocr)


def extract(data: bytes, mime_type: str) -> ExtractionResult:
    kind = SUPPORTED_MIME.get(mime_type)
    if kind == "pdf":
        return _extract_pdf(data)
    if kind == "docx":
        return _extract_docx(data)
    if kind == "txt":
        return ExtractionResult(pages=[ExtractedPage(None, data.decode("utf-8", errors="replace"))])
    raise ValueError(f"Unsupported mime type: {mime_type}")


def _extract_pdf(data: bytes) -> ExtractionResult:
    from pypdf import PdfReader

    settings = get_settings()
    reader = PdfReader(io.BytesIO(data))
    result = ExtractionResult()
    ocr_needed: list[int] = []

    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if len(text) < 20:  # likely a scanned page with no usable text layer
            ocr_needed.append(i)
            result.pages.append(ExtractedPage(i, ""))
        else:
            result.pages.append(ExtractedPage(i, text))

    if ocr_needed and settings.ocr_enabled:
        _ocr_pages(data, ocr_needed, result)
    return result


def _ocr_pages(pdf_bytes: bytes, page_numbers: list[int], result: ExtractionResult) -> None:
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError:
        logger.warning("OCR requested but pytesseract/pdf2image unavailable")
        return

    settings = get_settings()
    for page_no in page_numbers:
        try:
            images = convert_from_bytes(
                pdf_bytes, first_page=page_no, last_page=page_no, dpi=300
            )
            if images:
                text = pytesseract.image_to_string(images[0], lang=settings.ocr_languages)
                result.pages[page_no - 1] = ExtractedPage(page_no, text.strip(), used_ocr=True)
        except Exception:
            logger.exception(f"OCR failed for page {page_no}")


def _extract_docx(data: bytes) -> ExtractionResult:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for para in document.paragraphs:
        if not para.text.strip():
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            parts.append(f"\n## {para.text.strip()}\n")
        else:
            parts.append(para.text.strip())
    for table in document.tables:  # keep tables whole, row-per-line
        rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
        parts.append("\n".join(rows))
    return ExtractionResult(pages=[ExtractedPage(None, "\n\n".join(parts))])
